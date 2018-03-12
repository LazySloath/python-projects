# -*- coding: utf-8 -*-
"""
Created on Tue Feb 27 19:54:41 2018

@author: User

This bot scans for recent [IP] prompts by Syraphia and outputs it in a
Word Document
"""

import praw
import docx
from docx.shared import Inches
import bs4, requests
from PIL import Image
from io import BytesIO
import os

# Code for dealing with reddit
reddit = praw.Reddit(user_agent = 'WP/IP (by /u/LazySloath)', 
                    client_id = 'eRoARvm2LMffWA',
                    client_secret = 'f-NNksmavl-Lnd8dS7rTq8cBHOk')

syraphia = reddit.redditor('Syraphia')

posts = syraphia.submissions.new()
ip_link = {}

def get_link(post):
    """Get deviantart link from reddit post"""
    start = post.selftext.index('http')
    end = post.selftext[start:].index(')') + start
    return post.selftext[start:end]

def soupify(link):
    """Create BeautifulSoup object from a link"""
    
    res = requests.get(link)
    res.raise_for_status
    return bs4.BeautifulSoup(res.text, 'lxml')

def get_image(link):
    """Get image (in bytes) from a link"""
    soup = soupify(link)
    hit = soup.select('.dev-content-normal')[0]
    r = requests.get(hit.get('src'))
    return BytesIO(r.content)

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

if __name__ == '__main__':
    # Find 10 recent [IP] posts by Syraphia
    for post in posts:
        if '[IP]' in post.title:
            ip_link[post.title] = (get_link(post), post.url)
            if len(ip_link.keys()) == 10:
                break
    # Initialise document
    doc = docx.Document()
    for k, v in ip_link.items():
        p = doc.add_paragraph()
        # Add bold prompt title
        p.add_run(k).bold = True
        # Add hyperlinks
        for link in v:
            p = doc.add_paragraph()
            hyperlink = add_hyperlink(p, link, link, '0000EE', True)
        # Add picture
        doc.add_picture(get_image(v[0]), width = Inches(5))
        doc.add_page_break()
        print('{}\n{}\n{}\n'.format(k, v[0], v[1]))
    # Save and open document   
    doc.save('wpipviewer.docx')
    os.system("start " + 'wpipviewer.docx')
    print('ok done!')
